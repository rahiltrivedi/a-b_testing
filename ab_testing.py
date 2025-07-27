import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from statsmodels.stats.proportion import proportions_ztest
from docx import Document
from docx.shared import Inches
import os

# Optional PDF conversion (works on Windows/macOS only)
try:
    from docx2pdf import convert
    pdf_export = True
except ImportError:
    pdf_export = False
    print("⚠️ 'docx2pdf' not installed. PDF export will be skipped.")

# Create folders
os.makedirs("report", exist_ok=True)
os.makedirs("plots", exist_ok=True)

# Load dataset
ab_data = pd.read_csv("ab_test_data.csv")

# Group sizes and conversion rates
group_sizes = ab_data['group'].value_counts()
overall_conversion = ab_data['converted'].mean() * 100
conversion_by_group = ab_data.groupby('group')['converted'].mean() * 100

# Plot conversion rates
sns.barplot(x='group', y='converted', data=ab_data, ci=None)
plt.title("Conversion Rate by Group")
plt.ylabel("Conversion Rate")
plt.ylim(0, 0.2)
plot_path = "plots/conversion_rate_plot.png"
plt.savefig(plot_path)
plt.close()

# Hypothesis test
converted = ab_data.groupby('group')['converted'].sum()
n_obs = ab_data['group'].value_counts().sort_index()
z_stat, p_val = proportions_ztest(count=converted, nobs=n_obs)

alpha = 0.05
conclusion = (
    "✅ Statistically significant — the variant likely improves conversion."
    if p_val < alpha else
    "❌ Not statistically significant — no strong evidence of difference."
)

# Prepare summary
report_text = f"""
A/B TESTING ANALYSIS REPORT

🔍 Group Size:
{group_sizes.to_string()}

🔍 Overall Conversion Rate: {overall_conversion:.2f}%

📊 Conversion Rate by Group:
{conversion_by_group.to_string()}

🧪 Hypothesis Test Results:
Z-statistic: {z_stat:.3f}
P-value: {p_val:.4f}

Conclusion:
{conclusion}
"""

# Generate Word document
doc = Document()
doc.add_heading("A/B Testing Report", level=1)
for line in report_text.strip().split('\n'):
    doc.add_paragraph(line.strip())

doc.add_paragraph("\n📈 Conversion Rate Plot:")
doc.add_picture(plot_path, width=Inches(5))

docx_path = "report/ab_test_report.docx"
doc.save(docx_path)
print(f"✅ Word report saved to: {docx_path}")

# Convert to PDF (if possible)
if pdf_export:
    pdf_path = "report/ab_test_report.pdf"
    convert(docx_path, pdf_path)
    print(f"PDF report saved to: {pdf_path}")
else:
    print("PDF report not generated (requires 'docx2pdf' and Windows/macOS).")
